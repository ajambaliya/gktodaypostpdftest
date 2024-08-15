import os
import asyncio
import requests
import logging
from telegram import Bot
from telegram.constants import ParseMode
from telegram.error import TelegramError
from pymongo import MongoClient
import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import subprocess
import tempfile
import math

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Read environment variables
mongo_uri = os.getenv('MONGO_CONNECTION_STRING')
BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
DEFAULT_CHANNEL = os.getenv('TELEGRAM_CHANNEL_USERNAME')
TEMPLATE_URL = os.getenv('TEMPLATE_URL', 'https://docs.google.com/document/d/12t9nJzPPHqXbRcH3As4PitcJi9w0SeuD/edit?usp=sharing&ouid=108520131839767724661&rtpof=true&sd=true')

# Debug logging for environment variables
logger.info(f"Bot Token: {BOT_TOKEN}")
logger.info(f"Telegram Channel: {DEFAULT_CHANNEL}")
logger.info(f"Mongo URI: {mongo_uri}")

# Initialize MongoDB client and Telegram bot
client = MongoClient(mongo_uri)
bot = Bot(token=BOT_TOKEN)

# Define functions for the bot
def fetch_collections(database_name):
    db = client[database_name]
    return db.list_collection_names()

def fetch_questions_from_collection(database_name, collection_name, num_questions):
    db = client[database_name]
    collection = db[collection_name]
    questions = collection.aggregate([{ '$sample': { 'size': num_questions } }])
    return list(questions)

def get_correct_option_index(answer_key):
    option_mapping = {'a': 0, 'b': 1, 'c': 2, 'd': 3}
    return option_mapping.get(answer_key.lower(), None)

def get_quiz_day():
    db = client['QuizDays']
    collection = db['Days']
    today = datetime.now().date()
    today_datetime = datetime.combine(today, datetime.min.time())

    day_record = collection.find_one({'date': today_datetime})
    
    if day_record:
        return day_record['day']
    else:
        last_day_record = collection.find_one(sort=[('date', -1)])
        new_day = 1 if not last_day_record else last_day_record['day'] + 1
        
        collection.insert_one({'date': today_datetime, 'day': new_day})
        return new_day

def get_quiz_number(collection_name):
    db = client['QuizCounters']
    collection = db['Counters']
    counter_record = collection.find_one({'collection_name': collection_name})
    
    if counter_record:
        new_count = counter_record['count'] + 1
        collection.update_one({'collection_name': collection_name}, {'$set': {'count': new_count}})
        return new_count
    else:
        collection.insert_one({'collection_name': collection_name, 'count': 1})
        return 1

def get_overall_quiz_number():
    db = client['QuizCounters']
    collection = db['OverallCounter']
    counter_record = collection.find_one({'counter_name': 'overall_quiz'})
    
    if counter_record:
        new_count = counter_record['count'] + 1
        collection.update_one({'counter_name': 'overall_quiz'}, {'$set': {'count': new_count}})
        return new_count
    else:
        collection.insert_one({'counter_name': 'overall_quiz', 'count': 1})
        return 1

async def send_intro_message(collection_name, num_questions, quiz_number, overall_quiz_number):
    day = get_quiz_day()
    intro_message = (
        f"🎯 *આજની કવિઝ - Day {day} - {collection_name} Quiz {quiz_number}* 🎯\n\n"
        f"📚 વિષય: *{collection_name}*\n"
        f"🔢 પ્રશ્નોની સંખ્યા: *{num_questions}*\n"
        f"🔢 કવિઝ નંબર: *{overall_quiz_number}*\n\n"
        f"🕐 અમારા ટેલીગ્રામ ચેનલમાં દરરોજ બપોરે *1 વાગ્યે* અને રાત્રે *9 વાગ્યે* "
        f"*{num_questions}* પ્રશ્નોની કવિઝ મુકવામાં આવે છે.\n\n"
        f"🔗 *Join* : @CurrentAdda\n\n"
        f"🏆 તૈયાર રહો! કવિઝ શરૂ થવાની તૈયારીમાં છે... 🚀"
    )
    
    try:
        await bot.send_message(
            chat_id=DEFAULT_CHANNEL,
            text=intro_message,
            parse_mode=ParseMode.MARKDOWN
        )
        logger.info("Intro message sent successfully")
    except TelegramError as e:
        logger.error(f"Error sending intro message: {e}")

async def send_quiz_to_channel(question, options, correct_option_index, explanation):
    question_text = f"{question}\n[@CurrentAdda]"
    
    if explanation is None or (isinstance(explanation, float) and math.isnan(explanation)):
        explanation = "@CurrentAdda"
    
    try:
        await bot.send_poll(
            chat_id=DEFAULT_CHANNEL,
            question=question_text,
            options=options,
            type='quiz',
            correct_option_id=correct_option_index,
            explanation=explanation,
            is_anonymous=True,
            allows_multiple_answers=False,
        )
        logger.info(f"Quiz sent successfully: {question}")
    except TelegramError as e:
        logger.error(f"Error sending quiz: {e}")

def download_template(url):
    download_url = url.replace('/edit?usp=sharing', '/export?format=docx')
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading template: {e}")
        raise

def update_document_with_content(doc_io, intro_message, questions, collection_name, quiz_number):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx_file:
        temp_docx_file.write(doc_io.read())
        temp_docx_path = temp_docx_file.name
    
    doc = Document(temp_docx_path)
    
    content_start = None
    content_end = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if '<<START_CONTENT>>' in paragraph.text:
            content_start = i
        elif '<<END_CONTENT>>' in paragraph.text:
            content_end = i
            break
    
    if content_start is not None and content_end is not None:
        # Clear existing paragraphs between placeholders
        for i in range(content_end - 1, content_start, -1):
            doc._element.body.remove(doc.paragraphs[i]._element)
        
        # Insert intro message
        intro_para = doc.paragraphs[content_start]
        intro_para.text = intro_message
        intro_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        intro_para.style.font.size = Pt(14)
        intro_para.style.font.bold = True
        intro_para.style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue color
        
        # Insert questions
        for q in questions:
            question_text = q.get('Question', 'No question text')
            question_paragraph = doc.add_paragraph(f"Q: {question_text}", style='Normal')
            question_paragraph.insert_paragraph_before()
            question_paragraph.style.font.size = Pt(12)
            
            options = [
                f"A) {q.get('Option A', 'No option')}",
                f"B) {q.get('Option B', 'No option')}",
                f"C) {q.get('Option C', 'No option')}",
                f"D) {q.get('Option D', 'No option')}"
            ]
            for option in options:
                option_paragraph = doc.add_paragraph(option, style='Normal')
                option_paragraph.style.font.size = Pt(10)
            
            answer = f"Answer: {q.get('Answer', 'Not provided')}"
            answer_paragraph = doc.add_paragraph(answer, style='Normal')
            answer_paragraph.style.font.size = Pt(10)
            answer_paragraph.style.font.bold = True
            
            doc.add_paragraph()  # Add a blank line between questions
    else:
        logger.warning("Could not find both <<START_CONTENT>> and <<END_CONTENT>> placeholders in the document.")

    updated_doc_path = os.path.join(tempfile.gettempdir(), f'{collection_name} Quiz {quiz_number}.docx')
    doc.save(updated_doc_path)
    
    return updated_doc_path

def convert_docx_to_pdf(docx_file, pdf_path):
    try:
        output_dir = os.path.dirname(pdf_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_file],
            check=True, capture_output=True, text=True
        )
        logger.info(f"LibreOffice conversion output: {result.stdout}")
        logger.error(f"LibreOffice conversion error output: {result.stderr}")
        
        pdf_temp_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_file))[0] + '.pdf')
        if os.path.exists(pdf_temp_path):
            os.rename(pdf_temp_path, pdf_path)
            logger.info(f"Successfully converted DOCX to PDF: {pdf_path}")
        else:
            raise FileNotFoundError(f"PDF file not found at expected location: {pdf_temp_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e}")
        logger.error(f"LibreOffice stderr: {e.stderr}")
        raise
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {e}")
        raise

async def send_pdf_to_channel(pdf_path, caption, collection_name, quiz_number, overall_quiz_number):
    attractive_caption = (
        f"🎉 *{collection_name} Quiz {quiz_number} (Overall Quiz {overall_quiz_number}) is now available!* 🎉\n\n"
        f"📚 Boost your knowledge with our latest quiz.\n"
        f"🧠 Challenge yourself and learn something new!\n\n"
        f"📥 Download the PDF and start quizzing.\n"
        f"🔗 Don't forget to join @CurrentAdda for daily updates!\n\n"
        f"#Quiz #{collection_name.replace(' ', '')} #Quiz{overall_quiz_number}"
    )
    
    try:
        with open(pdf_path, 'rb') as pdf_file:
            await bot.send_document(
                chat_id=DEFAULT_CHANNEL,
                document=pdf_file,
                caption=attractive_caption,
                parse_mode=ParseMode.MARKDOWN
            )
        logger.info(f"PDF sent successfully")
    except TelegramError as e:
        logger.error(f"Failed to send PDF: {e.message}")

async def main():
    collections = fetch_collections('MasterQuestions')
    
    if not collections:
        logger.info("No collections found in the database.")
        return
    
    selected_collection = random.choice(collections)
    num_questions = 3
    questions = fetch_questions_from_collection('MasterQuestions', selected_collection, num_questions)
    
    quiz_number = get_quiz_number(selected_collection)
    overall_quiz_number = get_overall_quiz_number()
    
    intro_message = (
        f"🎯 *Day {get_quiz_day()} - {selected_collection} Quiz {quiz_number}* 🎯\n\n"
        f"📚 *Quiz Collection*: {selected_collection}\n"
        f"🔢 *Number of Questions*: {num_questions}\n"
        f"🔢 *Overall Quiz Number*: {overall_quiz_number}\n\n"
        f"🕐 Daily quizzes are posted at *1 PM* and *9 PM*.\n\n"
        f"🔗 *Join*: @CurrentAdda\n\n"
        f"🏆 Get ready for your quiz! 🚀"
    )
    
    await send_intro_message(selected_collection, num_questions, quiz_number, overall_quiz_number)
    await asyncio.sleep(5)
    
    for question in questions:
        question_text = question.get('Question', 'No question text')
        options = [str(question.get('Option A', 'No option')), str(question.get('Option B', 'No option')), 
                   str(question.get('Option C', 'No option')), str(question.get('Option D', 'No option'))]
        correct_option_index = get_correct_option_index(question.get('Answer', 'a'))
        explanation = question.get('Explanation', None)
        
        if correct_option_index is not None:
            await send_quiz_to_channel(question_text, options, correct_option_index, explanation)
            await asyncio.sleep(3)
    
    template_io = download_template(TEMPLATE_URL)
    updated_doc_path = update_document_with_content(template_io, intro_message, questions, selected_collection, quiz_number)
    pdf_path = os.path.join(tempfile.gettempdir(), f'{selected_collection} Quiz {quiz_number} - Overall {overall_quiz_number}.pdf')
    
    convert_docx_to_pdf(updated_doc_path, pdf_path)
    await send_pdf_to_channel(pdf_path, f"{selected_collection} Quiz {quiz_number} - Overall {overall_quiz_number} - {datetime.now().strftime('%d %B %Y')}", selected_collection, quiz_number, overall_quiz_number)

if __name__ == "__main__":
    asyncio.run(main())
